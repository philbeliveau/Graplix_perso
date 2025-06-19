#!/usr/bin/env python3
"""
Enhanced LLM test with automatic OpenAI API key loading from .env file.
Usage: python enhanced_llm_test.py
"""

import os
import sys
import json
import base64
import traceback
import re
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
import time
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Configuration
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
if not OPENAI_API_KEY:
    print("❌ Error: OPENAI_API_KEY not found in .env file")
    print("Please add your OpenAI API key to the .env file:")
    print("OPENAI_API_KEY=your_key_here")
    sys.exit(1)
DOCUMENT_PASSWORD = "Hubert"  # Password for protected files

print(f"🔐 Using password '{DOCUMENT_PASSWORD}' for protected documents")
print(f"🔑 OpenAI API key configured")

try:
    import openai
    from PIL import Image
    print("✅ Core libraries loaded")
    
    # Try to import optional libraries
    try:
        from pdf2image import convert_from_path
        PDF_SUPPORT = True
        print("✅ PDF processing available")
    except ImportError:
        PDF_SUPPORT = False
        print("⚠️ PDF processing not available (install pdf2image)")
        
    try:
        import openpyxl
        EXCEL_SUPPORT = True
        print("✅ Excel processing available")
    except ImportError:
        EXCEL_SUPPORT = False
        print("⚠️ Excel processing not available (install openpyxl)")
        
    try:
        from docx import Document
        DOCX_SUPPORT = True
        print("✅ Word processing available")  
    except ImportError:
        DOCX_SUPPORT = False
        print("⚠️ Word processing not available (install python-docx)")
        
    # Tesseract OCR as fallback
    try:
        import pytesseract
        TESSERACT_SUPPORT = True
        print("✅ Tesseract OCR available as fallback")
    except ImportError:
        TESSERACT_SUPPORT = False
        print("⚠️ Tesseract OCR not available (install pytesseract)")
        
except ImportError as e:
    print(f"❌ Critical import error: {e}")
    print("Install: pip install openai pillow pdf2image python-docx openpyxl")
    sys.exit(1)

# Initialize OpenAI client
client = openai.OpenAI(api_key=OPENAI_API_KEY)

def extract_pii_entities(text: str) -> Dict[str, List[str]]:
    """Extract actual PII entities using precise regex patterns."""
    if not text or text.strip() == "I'm sorry, but I can't assist with that.":
        return {}
    
    found_pii = {}
    
    # Email addresses - this works well
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text, re.IGNORECASE)
    if emails:
        found_pii["emails"] = list(set(emails))
    
    # Phone numbers - this works well
    phone_patterns = [
        r'\b\(\d{3}\)\s*\d{3}[-.\s]?\d{4}\b',  # (514) 123-4567
        r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # 514-123-4567 or 514.123.4567
        r'\b\+1[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # +1-514-123-4567
        r'\b\(\d{3}\)\s?\d{3}\s?\d{4}\b'  # (514) 123 4567
    ]
    phones = []
    for pattern in phone_patterns:
        phones.extend(re.findall(pattern, text))
    if phones:
        found_pii["phone_numbers"] = list(set(phones))
    
    # MUCH MORE PRECISE Name patterns
    names = []
    
    # 1. Form field patterns - most reliable
    form_name_patterns = [
        r'(?:Nom[,:]|Name[,:])\s*([A-Z][A-ZÀ-Ÿ\s]+(?:[A-Z][a-zà-ÿ]+))',  # Nom: LASTNAME
        r'(?:Prénom[,:]|First\s*name[,:]|Firstname[,:])\s*([A-Z][a-zà-ÿ]+)',  # Prénom: Firstname
        r'(?:Nom,\s*Prénom[,:])\s*([A-Z][A-ZÀ-Ÿ]+),\s*([A-Z][a-zà-ÿ]+)',  # Nom, Prénom : LASTNAME, Firstname
    ]
    
    for pattern in form_name_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if isinstance(match, tuple):
                # Handle tuple matches like (LASTNAME, Firstname)
                full_name = f"{match[1].strip()} {match[0].strip()}"
                names.append(full_name)
            else:
                names.append(match.strip())
    
    # 2. Specific known patterns from your documents
    specific_patterns = [
        r'\b([A-Z][A-Z]+),\s*([A-Z][a-zà-ÿ]+)\b',  # TREMBLAY, Steve
        r'\b([A-Z][a-zà-ÿ]+)\s+([A-Z][A-Z]+)\b',   # Steve TREMBLAY
        r'([A-Z][a-zà-ÿ]+)\s+([A-Z][a-zà-ÿ]+)(?:\s+\|)',  # From Excel: Therrien | Julien
        r'\b([A-Z][a-zà-ÿ-]+)\s*\|\s*([A-Z][a-zà-ÿ-]+)\b',  # Excel: lastname | firstname
        r'anonymous\s*\|\s*([A-Z][a-zà-ÿ-]+)\s*\|\s*([A-Z][a-zà-ÿ-]+)',  # Excel: anonymous | lastname | firstname
    ]
    
    for pattern in specific_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if len(match) == 2:
                # For patterns with two groups, create proper name
                if match[0].isupper() and not match[1].isupper():
                    # LASTNAME, Firstname
                    full_name = f"{match[1]} {match[0]}"
                elif not match[0].isupper() and match[1].isupper():
                    # Firstname LASTNAME  
                    full_name = f"{match[0]} {match[1]}"
                else:
                    # Both same case, keep order
                    full_name = f"{match[0]} {match[1]}"
                names.append(full_name)
    
    # 3. Filter and validate names
    validated_names = []
    for name in names:
        name = name.strip()
        if is_valid_name(name):
            validated_names.append(name)
    
    if validated_names:
        found_pii["names"] = list(set(validated_names))
    
    # Dates - keep existing patterns
    date_patterns = [
        r'\b\d{4}-\d{2}-\d{2}\b',  # YYYY-MM-DD
        r'\b\d{1,2}/\d{1,2}/\d{4}\b',  # DD/MM/YYYY or MM/DD/YYYY
        r'\b\d{1,2}-\d{1,2}-\d{4}\b',  # DD-MM-YYYY
        r'\b\d{1,2}\s+(?:janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+\d{4}\b',
        r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b'
    ]
    dates = []
    for pattern in date_patterns:
        dates.extend(re.findall(pattern, text, re.IGNORECASE))
    if dates:
        found_pii["dates"] = list(set(dates))
    
    # Addresses - more precise
    address_patterns = [
        r'\d+[,]?\s+(?:rue|street|boulevard|avenue|ave|blvd|road|rd)\s+[A-Za-záàâäçéèêëïîôùûüÿñ\s]+,?\s*[A-Za-z]{2,}',
        r'(?:Adresse|Address)\s*:\s*([^\n]+)',
        r'\d+\s+[A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*'  # Street addresses
    ]
    addresses = []
    for pattern in address_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        addresses.extend(matches)
    if addresses:
        found_pii["addresses"] = list(set([addr.strip() for addr in addresses if len(addr.strip()) > 10]))
    
    # Employee IDs and account numbers
    id_patterns = [
        r'(?:Employee\s+ID|ID\s+Employé[e]?)\s*:\s*([A-Z0-9]+)',
        r'(?:Numéro\s+de\s+certificat|Certificate\s+Number)\s*:\s*(\d+)',
        r'\b[A-Z]{3}\d{5,}\b'  # Pattern like EMP12345
    ]
    ids = []
    for pattern in id_patterns:
        ids.extend(re.findall(pattern, text, re.IGNORECASE))
    if ids:
        found_pii["employee_ids"] = list(set(ids))
    
    # Social Insurance Numbers (SIN) - Canadian format
    sin_pattern = r'\b\d{3}\s+\d{3}\s+\d{3}\b'
    sins = re.findall(sin_pattern, text)
    if sins:
        found_pii["social_insurance_numbers"] = list(set(sins))
    
    # Postal codes - keep existing
    postal_patterns = [
        r'\b[A-Z]\d[A-Z]\s*\d[A-Z]\d\b',  # Canadian postal codes
        r'\b\d{5}(?:-\d{4})?\b'  # US ZIP codes
    ]
    postal_codes = []
    for pattern in postal_patterns:
        postal_codes.extend(re.findall(pattern, text))
    if postal_codes:
        found_pii["postal_codes"] = list(set(postal_codes))
    
    return found_pii

def is_valid_name(name: str) -> bool:
    """Validate if a string is actually a name and not random text."""
    if not name or len(name.strip()) < 2:
        return False
    
    name = name.strip()
    
    # Remove obvious non-names
    invalid_indicators = [
        'date', 'naissance', 'birth', 'adresse', 'address', 'téléphone', 'phone',
        'email', 'courriel', 'poste', 'titre', 'département', 'service', 'division',
        'employé', 'employee', 'médecin', 'doctor', 'signature', 'horaire', 'travail',
        'work', 'bureau', 'office', 'province', 'canada', 'québec', 'ontario',
        'montreal', 'toronto', 'laval', 'gatineau', 'social', 'assurance', 'numéro',
        'number', 'compte', 'account', 'montréal', 'siège', 'formulaire', 'form',
        'clinique', 'hôpital', 'hospital', 'centre', 'center', 'université',
        'university', 'institut', 'institute', 'entreprise', 'company', 'groupe',
        'group', 'corporation', 'ltée', 'inc', 'pharmacien', 'ordonnance',
        'prescription', 'patient', 'dossier', 'file', 'document'
    ]
    
    name_lower = name.lower()
    for indicator in invalid_indicators:
        if indicator in name_lower:
            return False
    
    # Must have at least one letter
    if not re.search(r'[A-Za-zÀ-ÿ]', name):
        return False
    
    # Remove names that are too long (likely phrases)
    if len(name.split()) > 4:
        return False
    
    # Remove names with too many special characters
    if len(re.findall(r'[^\w\s\-àâäéèêëïîôùûüÿç]', name)) > 1:
        return False
    
    return True

def parse_openai_pii_response(response_text: str) -> Dict[str, List[str]]:
    """Parse OpenAI's structured response to extract PII information."""
    found_pii = {}
    
    if not response_text or "I'm sorry" in response_text or "can't assist" in response_text:
        return found_pii
    
    # Look for the PERSONAL_INFO section
    if "PERSONAL_INFO:" in response_text:
        pii_section = response_text.split("PERSONAL_INFO:")[1]
        
        # Parse each category
        categories = {
            "names": r"Names:\s*(.+?)(?:\n[A-Z][a-z]+:|$)",
            "emails": r"Emails:\s*(.+?)(?:\n[A-Z][a-z]+:|$)",
            "phones": r"Phones:\s*(.+?)(?:\n[A-Z][a-z]+:|$)",
            "dates": r"Dates:\s*(.+?)(?:\n[A-Z][a-z]+:|$)",
            "addresses": r"Addresses:\s*(.+?)(?:\n[A-Z][a-z]+:|$)",
            "employee_ids": r"IDs:\s*(.+?)(?:\n[A-Z][a-z]+:|$)",
            "other": r"Other:\s*(.+?)(?:\n[A-Z][a-z]+:|$)"
        }
        
        for category, pattern in categories.items():
            matches = re.findall(pattern, pii_section, re.DOTALL | re.IGNORECASE)
            if matches:
                # Clean up the extracted text
                items = []
                for match in matches:
                    # Split by common separators and clean
                    parts = re.split(r'[,;\n]+', match.strip())
                    for part in parts:
                        part = part.strip()
                        if part and part != "None" and part != "N/A" and part != "-" and len(part) > 1:
                            # Remove brackets if present
                            part = re.sub(r'^\[|\]$', '', part)
                            items.append(part)
                
                if items:
                    found_pii[category] = list(set(items))
    
    # Fallback: if no structured response, try basic regex on the text
    if not found_pii:
        # Use our existing function as fallback
        found_pii = extract_pii_entities(response_text)
    
    return found_pii

def encode_image_to_base64(image_input):
    """Convert image (path or PIL) to base64."""
    try:
        if isinstance(image_input, (str, Path)):
            image = Image.open(image_input)
        else:
            image = image_input
            
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Resize if too large for API
        if max(image.size) > 2048:
            image.thumbnail((2048, 2048), Image.Resampling.LANCZOS)
        
        buffer = BytesIO()
        image.save(buffer, format='PNG')
        return base64.b64encode(buffer.getvalue()).decode()
    except Exception as e:
        print(f"      ❌ Image encoding error: {e}")
        return None

def process_pdf_with_password(pdf_path):
    """Process PDF, trying with password first."""
    if not PDF_SUPPORT:
        return None, "PDF support not available"
    
    try:
        print(f"      🔐 Trying with password '{DOCUMENT_PASSWORD}'...")
        pages = convert_from_path(
            str(pdf_path),
            dpi=200,
            fmt='RGB',
            userpw=DOCUMENT_PASSWORD,
            first_page=1,
            last_page=3  # Limit to first 3 pages
        )
        print(f"      ✅ PDF unlocked with password - {len(pages)} pages")
        return pages, "unlocked_with_password"
        
    except Exception as pwd_error:
        print(f"      ⚠️ Password failed, trying without: {pwd_error}")
        try:
            pages = convert_from_path(
                str(pdf_path),
                dpi=200,
                fmt='RGB',
                first_page=1,
                last_page=3
            )
            print(f"      ✅ PDF opened without password - {len(pages)} pages")
            return pages, "no_password_needed"
            
        except Exception as no_pwd_error:
            print(f"      ❌ PDF failed both ways: {no_pwd_error}")
            return None, f"Failed: {no_pwd_error}"

def process_office_document(file_path):
    """Process Word/Excel documents with password protection support."""
    file_ext = file_path.suffix.lower()
    
    if file_ext in ['.doc', '.docx'] and DOCX_SUPPORT:
        try:
            print(f"      📄 Processing Word document...")
            
            # First try with password protection handling
            try:
                import msoffcrypto
                
                # Try to decrypt if password protected
                with open(file_path, 'rb') as f:
                    office_file = msoffcrypto.OfficeFile(f)
                    
                    if office_file.is_encrypted():
                        print(f"      🔐 Document is encrypted, trying password '{DOCUMENT_PASSWORD}'...")
                        office_file.load_key(password=DOCUMENT_PASSWORD)
                        
                        # Decrypt to temporary file
                        import tempfile
                        with tempfile.NamedTemporaryFile(suffix='.docx') as temp_file:
                            office_file.save(temp_file)
                            temp_file.seek(0)
                            doc = Document(temp_file.name)
                    else:
                        doc = Document(file_path)
                        
            except ImportError:
                print(f"      ⚠️ msoffcrypto not available, trying direct access...")
                doc = Document(file_path)
            except Exception as decrypt_error:
                print(f"      ⚠️ Decryption failed: {decrypt_error}, trying direct access...")
                doc = Document(file_path)
            
            # Extract text content
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            full_text = '\n'.join(text_content)
            return {
                "success": True,
                "text_content": full_text,
                "text_length": len(full_text),
                "method": "password_aware_extraction",
                "extracted_pii": extract_pii_entities(full_text)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Word processing failed: {e}",
                "note": "Document may be corrupted or use unsupported encryption"
            }
    
    elif file_ext in ['.xlsx', '.xls'] and EXCEL_SUPPORT:
        try:
            print(f"      📊 Processing Excel document...")
            
            # Try with password protection handling
            try:
                import msoffcrypto
                
                with open(file_path, 'rb') as f:
                    office_file = msoffcrypto.OfficeFile(f)
                    
                    if office_file.is_encrypted():
                        print(f"      🔐 Excel is encrypted, trying password '{DOCUMENT_PASSWORD}'...")
                        office_file.load_key(password=DOCUMENT_PASSWORD)
                        
                        import tempfile
                        with tempfile.NamedTemporaryFile(suffix='.xlsx') as temp_file:
                            office_file.save(temp_file)
                            temp_file.seek(0)
                            wb = openpyxl.load_workbook(temp_file.name, data_only=True)
                    else:
                        wb = openpyxl.load_workbook(file_path, data_only=True)
                        
            except ImportError:
                print(f"      ⚠️ msoffcrypto not available, trying direct access...")
                wb = openpyxl.load_workbook(file_path, data_only=True)
            except Exception as decrypt_error:
                print(f"      ⚠️ Decryption failed: {decrypt_error}, trying direct access...")
                wb = openpyxl.load_workbook(file_path, data_only=True)
            
            text_content = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_content.append(f"\n--- Sheet: {sheet_name} ---")
                
                # Extract data more intelligently for better PII detection
                for row_idx, row in enumerate(sheet.iter_rows(values_only=True), 1):
                    if any(cell is not None for cell in row):
                        # Clean and format row data
                        row_data = []
                        for cell in row:
                            if cell is not None:
                                cell_str = str(cell).strip()
                                if cell_str:
                                    row_data.append(cell_str)
                        
                        if row_data:
                            # For employee data, format as name-like patterns
                            if len(row_data) >= 2 and row_idx > 1:  # Skip header row
                                # Try to identify name patterns in Excel data
                                formatted_row = ' | '.join(row_data)
                                text_content.append(formatted_row)
                                
                                # Also add individual cells as potential names
                                for cell_data in row_data:
                                    if len(cell_data.split()) >= 2:  # Likely a full name
                                        text_content.append(f"Name: {cell_data}")
                            else:
                                text_content.append(' | '.join(row_data))
            
            full_text = '\n'.join(text_content)
            return {
                "success": True,
                "text_content": full_text,
                "text_length": len(full_text),
                "method": "password_aware_extraction",
                "extracted_pii": extract_pii_entities(full_text)
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Excel processing failed: {e}",
                "note": "Document may be corrupted or use unsupported encryption"
            }
    
    return {
        "success": False,
        "error": f"Unsupported office format: {file_ext}"
    }

def create_specialized_prompt(file_path):
    """Create simple prompts that work with OpenAI - just ask for text transcription."""
    file_name = file_path.name.lower()
    
    # The key insight: ask for transcription only, not PII extraction
    if "formulaire" in file_name or "form" in file_name:
        return "What text do you see in this French form? Please transcribe all the text exactly as written, including all form fields and values."
    
    elif "cv" in file_name or "curriculum" in file_name or "resume" in file_name:
        return "What text do you see in this resume? Please transcribe all the content exactly as written."
    
    elif "cheque" in file_name or "check" in file_name:
        return "What text and numbers do you see in this document? Please transcribe everything exactly as written."
    
    elif "consent" in file_name or "consentement" in file_name:
        return "What text do you see in this form? Please transcribe all the content exactly as written."
    
    else:
        return "What text do you see in this image? Please transcribe all visible text exactly as written."

def process_with_tesseract_fallback(image_input):
    """Fallback OCR using Tesseract when OpenAI fails."""
    if not TESSERACT_SUPPORT:
        return {"success": False, "error": "Tesseract not available"}
    
    try:
        if isinstance(image_input, str):
            # It's base64 data
            import base64
            image_data = base64.b64decode(image_input)
            image = Image.open(BytesIO(image_data))
        else:
            # It's a PIL image
            image = image_input
        
        # Preprocess image for better OCR
        import cv2
        import numpy as np
        
        # Convert PIL to numpy array
        img_array = np.array(image)
        
        # Convert to grayscale
        if len(img_array.shape) == 3:
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            gray = img_array
        
        # Apply threshold to make text clearer
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Convert back to PIL
        processed_image = Image.fromarray(thresh)
        
        # Extract text with Tesseract
        extracted_text = pytesseract.image_to_string(processed_image, lang='eng+fra')
        
        if not extracted_text.strip():
            return {"success": False, "error": "No text extracted"}
        
        # Extract PII from Tesseract result
        found_pii = extract_pii_entities(extracted_text)
        
        return {
            "success": True,
            "extracted_text": extracted_text,
            "text_length": len(extracted_text),
            "pii_analysis": found_pii,
            "total_pii_indicators": sum(len(items) for items in found_pii.values()),
            "pii_entities_found": len([item for items in found_pii.values() for item in items]),
            "method": "tesseract_ocr"
        }
        
    except Exception as e:
        return {
            "success": False,
            "error": f"Tesseract failed: {e}"
        }

def process_with_llm_ocr(file_path, image_data, model="gpt-4o"):
    """Process image with LLM OCR using specialized prompts, with Tesseract fallback."""
    
    prompt = create_specialized_prompt(file_path)
    
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{image_data}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=4000,
            temperature=0.0
        )
        
        extracted_text = response.choices[0].message.content
        
        # Check if OpenAI refused to process
        if "sorry" in extracted_text.lower() or "can't assist" in extracted_text.lower():
            print("      🔄 OpenAI refused, trying Tesseract fallback...")
            fallback_result = process_with_tesseract_fallback(image_data)
            if fallback_result["success"]:
                fallback_result["method"] = "tesseract_fallback"
                fallback_result["usage"] = {
                    "prompt_tokens": response.usage.prompt_tokens,
                    "completion_tokens": response.usage.completion_tokens,
                    "total_tokens": response.usage.total_tokens,
                    "estimated_cost": (response.usage.prompt_tokens * 0.00015 + 
                                     response.usage.completion_tokens * 0.0006) / 1000
                }
                return fallback_result
            else:
                # Return the OpenAI result even if it's a refusal, but ensure usage is included
                pass
        
        # Extract PII from OpenAI's transcription using our improved regex
        found_pii = extract_pii_entities(extracted_text)
        
        return {
            "success": True,
            "extracted_text": extracted_text,
            "text_length": len(extracted_text),
            "pii_analysis": found_pii,
            "total_pii_indicators": sum(len(items) for items in found_pii.values()),
            "pii_entities_found": len([item for items in found_pii.values() for item in items]),
            "method": "openai_gpt4o",
            "usage": {
                "prompt_tokens": response.usage.prompt_tokens,
                "completion_tokens": response.usage.completion_tokens,
                "total_tokens": response.usage.total_tokens,
                "estimated_cost": (response.usage.prompt_tokens * 0.00015 + 
                                 response.usage.completion_tokens * 0.0006) / 1000
            }
        }
        
    except Exception as e:
        print(f"      ❌ OpenAI failed: {e}, trying Tesseract...")
        fallback_result = process_with_tesseract_fallback(image_data)
        if fallback_result["success"]:
            return fallback_result
        else:
            return {
                "success": False,
                "error": str(e),
                "traceback": traceback.format_exc()
            }

def test_all_files_with_passwords():
    """Test all files with proper password handling."""
    data_dir = Path("/Users/philippebeliveau/Desktop/Notebook/EZBI/GRAPLIX_GIT/data")
    
    if not data_dir.exists():
        print(f"❌ Data directory not found: {data_dir}")
        return
    
    print("\n🔍 Enhanced LLM OCR Test with Password Support")
    print("=" * 60)
    print(f"📁 Directory: {data_dir}")
    print(f"🔐 Password: '{DOCUMENT_PASSWORD}' (for protected files)")
    
    # Get all files (exclude temp files)
    all_files = [f for f in data_dir.iterdir() if f.is_file() and not f.name.startswith('~')]
    
    print(f"📊 Found {len(all_files)} files to process")
    
    results = []
    total_cost = 0
    successful_extractions = 0
    
    for i, file_path in enumerate(all_files, 1):
        print(f"\n📄 [{i}/{len(all_files)}] {file_path.name}")
        print(f"    Size: {file_path.stat().st_size / 1024:.1f} KB")
        
        file_result = {
            "file_name": file_path.name,
            "file_path": str(file_path),
            "file_type": file_path.suffix.lower(),
            "file_size_kb": round(file_path.stat().st_size / 1024, 1),
            "processing_results": {}
        }
        
        try:
            if file_path.suffix.lower() in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                # Process image files directly
                print("    🖼️ Processing as image file...")
                image_data = encode_image_to_base64(file_path)
                if image_data:
                    result = process_with_llm_ocr(file_path, image_data)
                    file_result["processing_results"]["llm_ocr"] = result
                    
                    if result.get("success"):
                        successful_extractions += 1
                        total_cost += result["usage"]["estimated_cost"]
                        pii_count = result.get("total_pii_indicators", 0)
                        print(f"    ✅ Success: {result['text_length']} chars, {pii_count} PII indicators")
                        
                        if result.get("pii_analysis"):
                            for category, items in result["pii_analysis"].items():
                                print(f"       {category}: {items[:3]}")  # Show first 3 items
                    else:
                        print(f"    ❌ Failed: {result.get('error', 'Unknown error')}")
                
            elif file_path.suffix.lower() == '.pdf':
                # Process PDF with password handling
                print("    📄 Processing PDF with password support...")
                pdf_pages, status = process_pdf_with_password(file_path)
                
                if pdf_pages:
                    pdf_results = []
                    for page_num, page in enumerate(pdf_pages, 1):
                        print(f"       Page {page_num}/{len(pdf_pages)}")
                        
                        page_data = encode_image_to_base64(page)
                        if page_data:
                            result = process_with_llm_ocr(file_path, page_data)
                            if result.get("success"):
                                result["page_number"] = page_num
                                pdf_results.append(result)
                                total_cost += result["usage"]["estimated_cost"]
                    
                    if pdf_results:
                        successful_extractions += 1
                        total_pii = sum(r.get("total_pii_indicators", 0) for r in pdf_results)
                        file_result["processing_results"]["pdf_pages"] = pdf_results
                        file_result["processing_results"]["pdf_status"] = status
                        print(f"    ✅ Success: {len(pdf_results)} pages, {total_pii} PII indicators")
                    else:
                        file_result["processing_results"]["error"] = "No pages processed successfully"
                        print(f"    ❌ No pages processed successfully")
                else:
                    file_result["processing_results"]["error"] = status
                    print(f"    ❌ PDF failed: {status}")
            
            elif file_path.suffix.lower() in ['.doc', '.docx', '.xlsx', '.xls']:
                # Process Office documents
                print(f"    📊 Processing {file_path.suffix.upper()} document...")
                office_result = process_office_document(file_path)
                file_result["processing_results"]["office_extraction"] = office_result
                
                if office_result.get("success"):
                    successful_extractions += 1
                    print(f"    ✅ Success: {office_result['text_length']} characters extracted")
                else:
                    print(f"    ❌ Failed: {office_result.get('error', 'Unknown error')}")
                    if office_result.get("note"):
                        print(f"    💡 Note: {office_result['note']}")
            
            elif file_path.suffix.lower() == '.txt':
                # Process text files directly
                print("    📝 Processing text file...")
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text_content = f.read()
                    
                    file_result["processing_results"]["direct_text"] = {
                        "success": True,
                        "text_length": len(text_content),
                        "preview": text_content[:300] + "..." if len(text_content) > 300 else text_content
                    }
                    successful_extractions += 1
                    print(f"    ✅ Success: {len(text_content)} characters")
                    
                except Exception as e:
                    file_result["processing_results"]["error"] = str(e)
                    print(f"    ❌ Text file error: {e}")
            
            else:
                file_result["processing_results"]["skipped"] = f"Unsupported format: {file_path.suffix}"
                print(f"    ⏭️ Skipped: {file_path.suffix} format not supported")
        
        except Exception as e:
            file_result["processing_results"]["error"] = str(e)
            print(f"    ❌ Processing error: {e}")
        
        results.append(file_result)
        time.sleep(0.5)  # Rate limiting
    
    # Save comprehensive results
    output_file = "enhanced_llm_test_results.json"
    summary = {
        "test_summary": {
            "total_files": len(all_files),
            "successful_extractions": successful_extractions,
            "success_rate": f"{(successful_extractions/len(all_files)*100):.1f}%",
            "total_cost_usd": round(total_cost, 4),
            "password_used": DOCUMENT_PASSWORD,
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
            "supported_formats": [".pdf", ".jpg", ".jpeg", ".png", ".doc", ".docx", ".xlsx", ".txt"]
        },
        "file_results": results
    }
    
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(summary, f, indent=2, ensure_ascii=False)
    
    # Print final summary
    print(f"\n🎯 FINAL RESULTS")
    print("=" * 30)
    print(f"📊 Files processed: {len(all_files)}")
    print(f"✅ Successful extractions: {successful_extractions}")
    print(f"📈 Success rate: {(successful_extractions/len(all_files)*100):.1f}%")
    print(f"💰 Total API cost: ${total_cost:.4f}")
    print(f"🔐 Password used: '{DOCUMENT_PASSWORD}'")
    print(f"📝 Results saved to: {output_file}")
    
    # Show file-by-file summary
    print(f"\n📋 File Summary:")
    for result in results:
        name = result["file_name"]
        proc_results = result["processing_results"]
        
        if "llm_ocr" in proc_results and proc_results["llm_ocr"].get("success"):
            pii = proc_results["llm_ocr"].get("total_pii_indicators", 0)
            print(f"✅ {name}: {pii} PII indicators (LLM OCR)")
        elif "pdf_pages" in proc_results:
            pages = len(proc_results["pdf_pages"])
            total_pii = sum(p.get("total_pii_indicators", 0) for p in proc_results["pdf_pages"])
            print(f"✅ {name}: {pages} pages, {total_pii} PII indicators (PDF)")
        elif "office_extraction" in proc_results and proc_results["office_extraction"].get("success"):
            chars = proc_results["office_extraction"]["text_length"]
            print(f"✅ {name}: {chars} characters (Office)")
        elif "direct_text" in proc_results:
            chars = proc_results["direct_text"]["text_length"]
            print(f"✅ {name}: {chars} characters (Text)")
        elif "skipped" in proc_results:
            print(f"⏭️ {name}: Skipped")
        else:
            print(f"❌ {name}: Failed")
    
    if successful_extractions >= len(all_files) * 0.7:
        print(f"\n🎉 EXCELLENT: {successful_extractions}/{len(all_files)} files processed successfully!")
        print("   LLM OCR is working well. Any dashboard issues are likely integration problems.")
    elif successful_extractions > 0:
        print(f"\n✅ PARTIAL SUCCESS: {successful_extractions}/{len(all_files)} files processed.")
        print("   Some files may need specific handling or format support.")
    else:
        print(f"\n⚠️ ISSUES DETECTED: Check API key, network, or file access problems.")

if __name__ == "__main__":
    test_all_files_with_passwords()
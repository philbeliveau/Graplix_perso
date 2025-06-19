"""Document processing utilities for various file formats."""

import io
import json
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

import cv2
import numpy as np
import pytesseract
from PIL import Image
from docx import Document
from PyPDF2 import PdfReader
from pdfminer.high_level import extract_text as pdf_extract_text

import sys
from pathlib import Path

# Add src to path for imports if needed
current_dir = Path(__file__).parent
src_dir = current_dir.parent
if str(src_dir) not in sys.path:
    sys.path.insert(0, str(src_dir))

from core.config import settings
from core.logging_config import get_logger

# Try to import LLM components, handle missing dependencies
try:
    from llm import llm_ocr_processor, OCRTaskType, LLM_PROCESSOR_AVAILABLE
except ImportError:
    llm_ocr_processor = None
    OCRTaskType = None
    LLM_PROCESSOR_AVAILABLE = False

# For password-protected documents
try:
    import msoffcrypto
    MSOFFCRYPTO_AVAILABLE = True
except ImportError:
    MSOFFCRYPTO_AVAILABLE = False

# For Excel files
try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# For EasyOCR
try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

# For PDF to image conversion
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

# Configure tesseract path for macOS Homebrew installation
pytesseract.pytesseract.tesseract_cmd = '/opt/homebrew/bin/tesseract'

logger = get_logger(__name__)


class DocumentProcessor:
    """Handles processing of various document formats."""
    
    def __init__(self):
        """Initialize the document processor."""
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,  # Limited support
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.txt': self._process_text,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image,
            '.tiff': self._process_image,
            '.bmp': self._process_image,
        }
        
        # Configure Tesseract
        pytesseract.pytesseract.tesseract_cmd = settings.processing.tesseract_cmd
        
        # Initialize EasyOCR reader if available
        self.easyocr_reader = None
        if EASYOCR_AVAILABLE and settings.processing.ocr_engine in ['easyocr', 'both']:
            try:
                # Initialize with English and French languages
                self.easyocr_reader = easyocr.Reader(['en', 'fr'], gpu=settings.processing.easyocr_use_gpu)
                logger.info("EasyOCR reader initialized successfully")
            except Exception as e:
                logger.warning(f"Failed to initialize EasyOCR: {e}")
                self.easyocr_reader = None
        
        # Initialize LLM OCR processor
        self.llm_ocr_enabled = getattr(settings.processing, 'enable_llm_ocr', False) and LLM_PROCESSOR_AVAILABLE
        if self.llm_ocr_enabled and llm_ocr_processor:
            try:
                # llm_ocr_processor is already initialized as a global instance
                logger.info("LLM OCR processor available")
            except Exception as e:
                logger.warning(f"LLM OCR processor initialization issue: {e}")
                self.llm_ocr_enabled = False
        elif getattr(settings.processing, 'enable_llm_ocr', False) and not LLM_PROCESSOR_AVAILABLE:
            logger.warning("LLM OCR requested but dependencies not available. Install: pip install openai anthropic google-generativeai")
    
    def process_document(self, file_path: Union[str, Path], password: Optional[str] = None) -> Dict:
        """
        Process a document and extract text and metadata.
        
        Args:
            file_path: Path to the document file
            password: Optional password for protected documents
            
        Returns:
            Dictionary containing extracted text, metadata, and processing info
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > settings.processing.max_file_size_mb:
            raise ValueError(f"File too large: {file_size_mb:.1f}MB > {settings.processing.max_file_size_mb}MB")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        logger.info(f"Processing document: {file_path}")
        
        try:
            # Pass password to processing methods that support it
            if file_extension in ['.docx', '.doc', '.pdf', '.xlsx', '.xls']:
                result = self.supported_formats[file_extension](file_path, password)
            else:
                result = self.supported_formats[file_extension](file_path)
            
            result.update({
                'file_path': str(file_path),
                'file_name': file_path.name,
                'file_size_mb': round(file_size_mb, 2),
                'file_type': file_extension,
                'processing_status': 'success'
            })
            
            logger.info(f"Successfully processed: {file_path}")
            return result
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {str(e)}")
            raise
    
    def _process_pdf(self, file_path: Path, password: Optional[str] = None) -> Dict:
        """Process PDF files."""
        result = {
            'raw_text': '',
            'ocr_text': '',
            'pages': [],
            'metadata': {},
            'bounding_boxes': []
        }
        
        try:
            # Try extracting text with PyPDF2 first
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                # Handle password-protected PDFs
                if pdf_reader.is_encrypted:
                    if password:
                        try:
                            pdf_reader.decrypt(password)
                            logger.info(f"Successfully decrypted PDF with provided password")
                        except Exception as e:
                            logger.error(f"Failed to decrypt PDF with password: {e}")
                            raise ValueError(f"Invalid password for encrypted PDF: {file_path}")
                    else:
                        raise ValueError(f"PDF is encrypted but no password provided: {file_path}")
                
                result['metadata'] = {
                    'num_pages': len(pdf_reader.pages),
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                    'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')) if pdf_reader.metadata else '',
                    'is_encrypted': pdf_reader.is_encrypted
                }
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    result['pages'].append({
                        'page_number': page_num + 1,
                        'text': page_text,
                        'char_count': len(page_text)
                    })
                    result['raw_text'] += page_text + '\n'
        
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {e}, trying pdfminer")
            
            # Fallback to pdfminer
            try:
                result['raw_text'] = pdf_extract_text(str(file_path))
            except Exception as e2:
                logger.warning(f"pdfminer extraction failed: {e2}")
        
        # If text extraction failed or returned very little, try OCR
        if len(result['raw_text'].strip()) < 50:
            logger.info("Text extraction yielded little content, attempting OCR")
            try:
                result['ocr_text'] = self._ocr_pdf(file_path, password)
            except Exception as e:
                logger.error(f"OCR failed: {e}")
        
        return result
    
    def _process_docx(self, file_path: Path, password: Optional[str] = None) -> Dict:
        """Process DOCX files."""
        result = {
            'raw_text': '',
            'ocr_text': '',
            'paragraphs': [],
            'metadata': {},
            'bounding_boxes': []
        }
        
        try:
            # Handle password-protected DOCX files and .doc format detection
            doc_to_process = file_path
            temp_file = None
            doc = None
            
            if password and MSOFFCRYPTO_AVAILABLE:
                try:
                    # Check if the file is encrypted
                    with open(file_path, 'rb') as f:
                        office_file = msoffcrypto.OfficeFile(f)
                        if office_file.is_encrypted():
                            # Create temporary decrypted file
                            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                            office_file.load_key(password=password)
                            office_file.decrypt(temp_file)
                            temp_file.close()
                            doc_to_process = temp_file.name
                            logger.info(f"Successfully decrypted DOCX with provided password")
                except Exception as e:
                    logger.error(f"Failed to decrypt DOCX with password: {e}")
                    if temp_file:
                        temp_file.close()
                    raise ValueError(f"Invalid password for encrypted DOCX: {file_path}")
            
            # Try to process the document
            try:
                doc = Document(doc_to_process)
            except Exception as e:
                # Special handling for .doc files (older format)
                if file_path.suffix.lower() == '.doc':
                    logger.warning(f"Old .doc format detected: {file_path}. Consider converting to .docx format.")
                    raise ValueError(f"Old .doc format not supported. Please convert {file_path.name} to .docx format.")
                else:
                    # Re-raise the original error for .docx files
                    raise
            
            # Extract metadata
            core_props = doc.core_properties
            result['metadata'] = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'num_paragraphs': len(doc.paragraphs),
                'is_encrypted': doc_to_process != str(file_path)
            }
            
            # Extract text from paragraphs
            for para_num, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text
                result['paragraphs'].append({
                    'paragraph_number': para_num + 1,
                    'text': para_text,
                    'char_count': len(para_text)
                })
                result['raw_text'] += para_text + '\n'
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        result['raw_text'] += cell.text + ' '
                result['raw_text'] += '\n'
            
            # Clean up temporary file if created
            if temp_file and doc_to_process != str(file_path):
                try:
                    import os
                    os.unlink(temp_file.name)
                except Exception as cleanup_e:
                    logger.warning(f"Failed to cleanup temporary file: {cleanup_e}")
        
        except Exception as e:
            # Clean up temporary file on error
            if temp_file and doc_to_process != str(file_path):
                try:
                    import os
                    os.unlink(temp_file.name)
                except Exception as cleanup_e:
                    logger.warning(f"Failed to cleanup temporary file on error: {cleanup_e}")
            
            logger.error(f"DOCX processing failed: {e}")
            raise
        
        return result
    
    def _process_excel(self, file_path: Path, password: Optional[str] = None) -> Dict:
        """Process Excel files (.xlsx, .xls)."""
        result = {
            'raw_text': '',
            'sheets': [],
            'metadata': {},
            'bounding_boxes': []
        }
        
        if not OPENPYXL_AVAILABLE:
            raise ValueError("openpyxl not available. Cannot process Excel files.")
        
        try:
            # Handle password-protected Excel files
            workbook_to_process = file_path
            temp_file = None
            
            if password and MSOFFCRYPTO_AVAILABLE:
                try:
                    # Check if the file is encrypted
                    with open(file_path, 'rb') as f:
                        office_file = msoffcrypto.OfficeFile(f)
                        if office_file.is_encrypted():
                            # Create temporary decrypted file
                            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
                            office_file.load_key(password=password)
                            office_file.decrypt(temp_file)
                            temp_file.close()
                            workbook_to_process = temp_file.name
                            logger.info(f"Successfully decrypted Excel with provided password")
                except Exception as e:
                    logger.error(f"Failed to decrypt Excel with password: {e}")
                    if temp_file:
                        temp_file.close()
                    raise ValueError(f"Invalid password for encrypted Excel: {file_path}")
            
            # Process the Excel file (either original or decrypted)
            workbook = openpyxl.load_workbook(workbook_to_process, data_only=True)
            
            # Extract metadata
            result['metadata'] = {
                'sheet_names': workbook.sheetnames,
                'num_sheets': len(workbook.sheetnames),
                'is_encrypted': workbook_to_process != str(file_path)
            }
            
            # Process each sheet
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = ''
                rows_data = []
                
                # Extract data from cells
                for row_num, row in enumerate(sheet.iter_rows(values_only=True), 1):
                    row_text_parts = []
                    for cell_value in row:
                        if cell_value is not None:
                            cell_text = str(cell_value)
                            row_text_parts.append(cell_text)
                            sheet_text += cell_text + ' '
                    
                    if row_text_parts:  # Only add non-empty rows
                        rows_data.append({
                            'row_number': row_num,
                            'cells': row_text_parts,
                            'text': ' '.join(row_text_parts)
                        })
                    
                    sheet_text += '\n'
                
                result['sheets'].append({
                    'sheet_name': sheet_name,
                    'text': sheet_text,
                    'rows': rows_data,
                    'num_rows': len(rows_data)
                })
                
                result['raw_text'] += f"Sheet: {sheet_name}\n{sheet_text}\n"
            
            # Clean up temporary file if created
            if temp_file and workbook_to_process != str(file_path):
                try:
                    import os
                    os.unlink(temp_file.name)
                except Exception as cleanup_e:
                    logger.warning(f"Failed to cleanup temporary file: {cleanup_e}")
        
        except Exception as e:
            # Clean up temporary file on error
            if temp_file and workbook_to_process != str(file_path):
                try:
                    import os
                    os.unlink(temp_file.name)
                except Exception as cleanup_e:
                    logger.warning(f"Failed to cleanup temporary file on error: {cleanup_e}")
            
            logger.error(f"Excel processing failed: {e}")
            raise
        
        return result
    
    def _process_text(self, file_path: Path) -> Dict:
        """Process plain text files."""
        result = {
            'raw_text': '',
            'ocr_text': '',
            'lines': [],
            'metadata': {},
            'bounding_boxes': []
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                result['raw_text'] = content
                result['ocr_text'] = content  # For text files, raw and OCR text are the same
                
                lines = content.split('\n')
                result['lines'] = [{'line_number': i+1, 'text': line} for i, line in enumerate(lines)]
                result['metadata'] = {
                    'num_lines': len(lines),
                    'char_count': len(content),
                    'word_count': len(content.split()),
                    'encoding': 'utf-8'
                }
        
        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            raise
        
        return result
    
    def _process_image(self, file_path: Path) -> Dict:
        """Process image files with OCR."""
        result = {
            'raw_text': '',
            'ocr_text': '',
            'metadata': {},
            'bounding_boxes': []
        }
        
        try:
            # Load and enhance image
            image = cv2.imread(str(file_path))
            if image is None:
                raise ValueError(f"Could not load image: {file_path}")
            
            # Get image metadata
            height, width = image.shape[:2]
            file_size_kb = round(file_path.stat().st_size / 1024, 2)
            is_jpg = file_path.suffix.lower() in ['.jpg', '.jpeg']
            
            result['metadata'] = {
                'width': width,
                'height': height,
                'channels': image.shape[2] if len(image.shape) == 3 else 1,
                'file_size_kb': file_size_kb,
                'is_jpg': is_jpg,
                'format': file_path.suffix.lower(),
                'compression_ratio': file_size_kb / ((width * height * 3) / 1024) if width > 0 and height > 0 else 0
            }
            
            # Log image quality information
            if is_jpg:
                compression_ratio = result['metadata']['compression_ratio']
                logger.info(f"JPG file detected: {width}x{height}, {file_size_kb}KB, compression ratio: {compression_ratio:.3f}")
                if compression_ratio < 0.1:
                    logger.warning("Very high compression detected - this may severely impact OCR quality")
                elif compression_ratio < 0.3:
                    logger.warning("High compression detected - OCR quality may be affected")
            
            # Store file path for OCR configuration
            self._current_file_path = file_path
            
            # Enhance image for better OCR (pass file path for format-specific processing)
            enhanced_image = self._enhance_image_for_ocr(image, file_path)
            
            # Perform OCR based on selected engine
            ocr_engine = getattr(settings.processing, 'ocr_engine', 'tesseract')
            enable_llm_ocr = getattr(settings.processing, 'enable_llm_ocr', False)
            
            if enable_llm_ocr and self.llm_ocr_enabled:
                # Try LLM OCR first for better quality
                try:
                    llm_result = self._perform_llm_ocr(enhanced_image, result.copy())
                    
                    # If LLM OCR confidence is high, use it
                    if llm_result.get('confidence', 0) >= 0.8:
                        result = llm_result
                        logger.info(f"Used LLM OCR with confidence: {llm_result.get('confidence', 'N/A')}")
                    else:
                        # Fall back to traditional OCR and compare
                        traditional_result = self._perform_traditional_ocr(enhanced_image, result.copy(), ocr_engine)
                        result = self._combine_llm_traditional_results(llm_result, traditional_result)
                        logger.info("Combined LLM and traditional OCR results")
                        
                except Exception as e:
                    logger.warning(f"LLM OCR failed, falling back to traditional: {e}")
                    result = self._perform_traditional_ocr(enhanced_image, result, ocr_engine)
            else:
                # Use traditional OCR engines
                result = self._perform_traditional_ocr(enhanced_image, result, ocr_engine)
            
            logger.info(f"OCR extracted text from {file_path} using {settings.processing.ocr_engine} engine")
        
        except Exception as e:
            logger.error(f"Image processing failed: {e}")
            raise
        
        return result
    
    def _enhance_image_for_ocr(self, image: np.ndarray, file_path: Path = None) -> np.ndarray:
        """Enhance image quality for better OCR results, with special handling for JPG artifacts."""
        # Convert to grayscale
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image.copy()
        
        # Check if this is a JPG file (lossy compression artifacts)
        is_jpg = file_path and file_path.suffix.lower() in ['.jpg', '.jpeg']
        
        if is_jpg:
            # Form-optimized preprocessing for JPG files
            logger.info("Applying form-optimized JPG preprocessing for better OCR")
            
            # Apply CLAHE for better contrast (important for forms)
            clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
            enhanced = clahe.apply(gray)
            
            # Gentle denoising to reduce JPG artifacts without losing text detail
            denoised = cv2.fastNlMeansDenoising(enhanced, h=10, templateWindowSize=7, searchWindowSize=21)
            
            # Use adaptive thresholding (better for forms with varying lighting)
            # This preserves filled text better than OTSU thresholding
            binary = cv2.adaptiveThreshold(
                denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 15, 8
            )
            
            # Very minimal morphological operation to avoid removing important text
            kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (1, 1))
            processed = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
            
        else:
            # Standard preprocessing for PNG/TIFF/other lossless formats
            # Apply gentle denoising only
            denoised = cv2.fastNlMeansDenoising(gray, h=10)
            
            # Use gentle adaptive thresholding
            binary = cv2.adaptiveThreshold(
                denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            
            # Very minimal morphological operation
            kernel = np.ones((1, 1), np.uint8)
            processed = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
        
        # Convert back to BGR for consistency
        return cv2.cvtColor(processed, cv2.COLOR_GRAY2BGR)
    
    def _ocr_pdf(self, file_path: Path, password: Optional[str] = None) -> str:
        """Perform OCR on PDF pages by converting to images."""
        if not PDF2IMAGE_AVAILABLE:
            logger.warning("PDF OCR requires pdf2image library: pip install pdf2image")
            return ""
        
        try:
            logger.info(f"Converting PDF to images for OCR: {file_path}")
            
            # Convert PDF pages to images
            kwargs = {
                'pdf_path': str(file_path),
                'dpi': 200,  # Good balance of quality and speed
                'fmt': 'RGB',
                'thread_count': 1
            }
            
            # Add password if provided
            if password:
                kwargs['userpw'] = password
            
            # Convert pages to images
            pages = convert_from_path(**kwargs)
            
            if not pages:
                logger.warning("No pages found in PDF")
                return ""
            
            logger.info(f"Successfully converted {len(pages)} PDF pages to images")
            
            # Process each page with OCR
            all_text = []
            total_cost = 0
            
            for page_num, page_image in enumerate(pages, 1):
                try:
                    logger.info(f"Processing PDF page {page_num}/{len(pages)}")
                    
                    # Convert PIL image to numpy array for processing
                    page_array = np.array(page_image)
                    if len(page_array.shape) == 3:
                        # Convert RGB to BGR for OpenCV
                        page_array = cv2.cvtColor(page_array, cv2.COLOR_RGB2BGR)
                    
                    # Determine best OCR method
                    ocr_engine = getattr(settings.processing, 'ocr_engine', 'tesseract')
                    enable_llm_ocr = getattr(settings.processing, 'enable_llm_ocr', False)
                    
                    page_text = ""
                    
                    if enable_llm_ocr and self.llm_ocr_enabled:
                        # Try LLM OCR for this page
                        try:
                            logger.info(f"Using LLM OCR for PDF page {page_num}")
                            
                            # Create a temporary result dict for this page
                            page_result = {
                                'raw_text': '',
                                'ocr_text': '',
                                'metadata': {},
                                'bounding_boxes': []
                            }
                            
                            llm_result = self._perform_llm_ocr(page_array, page_result)
                            page_text = llm_result.get('ocr_text', '')
                            
                            # Track cost
                            llm_metadata = llm_result.get('llm_metadata', {})
                            if llm_metadata:
                                cost_info = llm_metadata.get('cost_info', {})
                                page_cost = cost_info.get('actual_cost', 0)
                                total_cost += page_cost
                                logger.info(f"Page {page_num} LLM OCR cost: ${page_cost:.6f}")
                            
                            if page_text.strip():
                                logger.info(f"LLM OCR extracted {len(page_text)} characters from page {page_num}")
                            else:
                                logger.warning(f"LLM OCR returned no text for page {page_num}, trying traditional OCR")
                                raise Exception("No text from LLM OCR")
                                
                        except Exception as llm_error:
                            logger.warning(f"LLM OCR failed for page {page_num}: {llm_error}, falling back to traditional OCR")
                            page_text = self._perform_traditional_ocr_on_image(page_array)
                    else:
                        # Use traditional OCR
                        page_text = self._perform_traditional_ocr_on_image(page_array)
                    
                    if page_text.strip():
                        all_text.append(f"--- Page {page_num} ---\n{page_text}")
                        logger.info(f"Successfully extracted text from PDF page {page_num}")
                    else:
                        logger.warning(f"No text extracted from PDF page {page_num}")
                    
                except Exception as page_error:
                    logger.error(f"Error processing PDF page {page_num}: {page_error}")
                    all_text.append(f"--- Page {page_num} ---\n[Error processing page: {page_error}]")
            
            final_text = "\n\n".join(all_text)
            
            if total_cost > 0:
                logger.info(f"Total LLM OCR cost for PDF: ${total_cost:.6f}")
            
            logger.info(f"PDF OCR completed. Extracted {len(final_text)} total characters from {len(pages)} pages")
            return final_text
            
        except Exception as e:
            logger.error(f"PDF OCR failed: {e}")
            return ""
    
    def _perform_traditional_ocr_on_image(self, image_array: np.ndarray) -> str:
        """Perform traditional OCR on a numpy image array."""
        try:
            # Enhance image for OCR
            enhanced_image = self._enhance_image_for_ocr(image_array)
            
            # Create temporary result dict
            temp_result = {
                'raw_text': '',
                'ocr_text': '',
                'metadata': {},
                'bounding_boxes': []
            }
            
            # Use the appropriate traditional OCR method
            ocr_engine = getattr(settings.processing, 'ocr_engine', 'tesseract')
            
            if ocr_engine == 'easyocr' and self.easyocr_reader:
                result = self._perform_easyocr(enhanced_image, temp_result)
            elif ocr_engine == 'both' and self.easyocr_reader:
                tesseract_result = self._perform_tesseract_ocr(enhanced_image, temp_result.copy())
                easyocr_result = self._perform_easyocr(enhanced_image, temp_result.copy())
                result = self._combine_ocr_results(tesseract_result, easyocr_result)
            else:
                result = self._perform_tesseract_ocr(enhanced_image, temp_result)
            
            return result.get('ocr_text', '')
            
        except Exception as e:
            logger.error(f"Traditional OCR on image failed: {e}")
            return ""
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self.supported_formats.keys())
    
    def is_supported_format(self, file_path: Union[str, Path]) -> bool:
        """Check if file format is supported."""
        file_path = Path(file_path)
        return file_path.suffix.lower() in self.supported_formats
    
    def _perform_tesseract_ocr(self, image: np.ndarray, result: Dict) -> Dict:
        """Perform OCR using Tesseract."""
        # Convert to PIL Image
        pil_image = Image.fromarray(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
        
        # Extract text with bounding box information using optimized OCR config
        # Try multiple OCR configurations for better results
        configs_to_try = []
        
        if hasattr(self, '_current_file_path') and self._current_file_path:
            is_jpg = self._current_file_path.suffix.lower() in ['.jpg', '.jpeg']
            if is_jpg:
                # Multiple configurations for JPG files (try in order of preference)
                # Forms often need PSM 6 (uniform text block) for best results
                configs_to_try = [
                    ('--oem 3 --psm 6', 'JPG forms - uniform text block'),
                    ('--oem 3 --psm 4', 'JPG forms - single column'),
                    ('--oem 3 --psm 11', 'JPG forms - sparse text'),
                    ('--oem 3 --psm 13', 'JPG forms - raw line')
                ]
            else:
                # Standard configurations for clean images
                configs_to_try = [
                    ('--oem 3 --psm 4', 'Standard - single column'),
                    ('--oem 3 --psm 6', 'Standard - uniform text block')
                ]
        else:
            configs_to_try = [('--oem 3 --psm 4', 'Default configuration')]
        
        # Try each configuration until we get reasonable results
        best_result = None
        best_confidence = 0
        
        for ocr_config, config_name in configs_to_try:
            try:
                logger.debug(f"Trying OCR config: {config_name} ({ocr_config})")
                ocr_data = pytesseract.image_to_data(
                    pil_image,
                    lang=settings.processing.ocr_languages,
                    config=ocr_config,
                    output_type=pytesseract.Output.DICT
                )
                
                # Check quality of this result
                text_parts = []
                total_confidence = 0
                valid_detections = 0
                
                for i in range(len(ocr_data['text'])):
                    text = ocr_data['text'][i].strip()
                    confidence = int(ocr_data['conf'][i])
                    
                    if text and confidence > 30:
                        text_parts.append(text)
                        total_confidence += confidence
                        valid_detections += 1
                
                # Calculate average confidence for this configuration
                avg_confidence = total_confidence / valid_detections if valid_detections > 0 else 0
                text_length = len(' '.join(text_parts))
                
                # Calculate quality score that prioritizes meaningful content over template lines
                full_text = ' '.join(text_parts)
                meaningful_content_score = 0
                
                # Boost score for actual data vs template lines
                if '___' not in full_text:  # No template underscores
                    meaningful_content_score += 20
                
                # Look for form data indicators
                form_indicators = ['nom', 'prénom', 'adresse', 'naissance', 'téléphone', 'date']
                for indicator in form_indicators:
                    if indicator.lower() in full_text.lower():
                        meaningful_content_score += 10
                
                # Look for actual filled data patterns (names, dates, addresses)
                data_patterns = [
                    r'\b[A-Z][a-z]+,?\s+[A-Z][a-z]+\b',  # Names like "TREMBLAY, Steve"
                    r'\b\d{4}-\d{2}-\d{2}\b',            # Dates like "1991-03-17"
                    r'\b\d{3}[-\s]\d{3}[-\s]\d{4}\b',    # Phone numbers
                    r'\b\d+\s+[A-Z][a-z]+\s+[A-Z][a-z]+\b'  # Addresses like "211 Avenue"
                ]
                
                import re
                for pattern in data_patterns:
                    if re.search(pattern, full_text):
                        meaningful_content_score += 15
                
                # Calculate final quality score
                quality_score = avg_confidence + meaningful_content_score
                
                logger.debug(f"Config {config_name}: {valid_detections} detections, confidence: {avg_confidence:.1f}, quality: {quality_score:.1f}, length: {text_length}")
                
                # Use this result if it has better quality score
                current_best_quality = best_result.get('quality_score', 0) if best_result else 0
                if quality_score > current_best_quality:
                    best_confidence = avg_confidence
                    best_result = {
                        'ocr_data': ocr_data,
                        'text_parts': text_parts,
                        'avg_confidence': avg_confidence,
                        'quality_score': quality_score,
                        'config_used': config_name
                    }
                
                # If we got excellent results, don't try more configs
                if quality_score > 100 and text_length > 20:
                    logger.info(f"Excellent OCR result with {config_name} (quality: {quality_score:.1f})")
                    break
                    
            except Exception as e:
                logger.warning(f"OCR config {config_name} failed: {e}")
                continue
        
        # Use the best result we found
        if best_result:
            ocr_data = best_result['ocr_data']
            text_parts = best_result['text_parts']
            logger.info(f"Using OCR result from: {best_result['config_used']} (confidence: {best_result['avg_confidence']:.1f})")
        else:
            # Fallback to basic OCR if all configurations failed
            logger.warning("All OCR configurations failed, using basic fallback")
            ocr_data = pytesseract.image_to_data(
                pil_image,
                lang=settings.processing.ocr_languages,
                output_type=pytesseract.Output.DICT
            )
            text_parts = []
        
        # Process final OCR results (if we didn't already process them above)
        if not best_result:
            text_parts = []
            for i in range(len(ocr_data['text'])):
                text = ocr_data['text'][i].strip()
                confidence = int(ocr_data['conf'][i])
                
                if text and confidence > 30:  # Filter low-confidence results
                    text_parts.append(text)
        
        # Create bounding boxes from final OCR data
        bounding_boxes = []
        for i in range(len(ocr_data['text'])):
            text = ocr_data['text'][i].strip()
            confidence = int(ocr_data['conf'][i])
            
            if text and confidence > 30:  # Filter low-confidence results
                # Store bounding box information
                bounding_boxes.append({
                    'text': text,
                    'confidence': confidence,
                    'x': ocr_data['left'][i],
                    'y': ocr_data['top'][i],
                    'width': ocr_data['width'][i],
                    'height': ocr_data['height'][i],
                    'engine': 'tesseract'
                })
        
        result['ocr_text'] = ' '.join(text_parts)
        result['raw_text'] = result['ocr_text']
        result['bounding_boxes'] = bounding_boxes
        result['ocr_engine'] = 'tesseract'
        
        return result
    
    def _perform_easyocr(self, image: np.ndarray, result: Dict) -> Dict:
        """Perform OCR using EasyOCR."""
        if not self.easyocr_reader:
            logger.warning("EasyOCR reader not available, falling back to Tesseract")
            return self._perform_tesseract_ocr(image, result)
        
        try:
            # Convert BGR to RGB for EasyOCR
            rgb_image = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
            
            # Perform OCR
            easyocr_results = self.easyocr_reader.readtext(rgb_image)
            
            # Process EasyOCR results
            text_parts = []
            bounding_boxes = []
            
            for detection in easyocr_results:
                bbox, text, confidence = detection
                
                if confidence > 0.3:  # Filter low-confidence results (0.3 = 30%)
                    text_parts.append(text)
                    
                    # EasyOCR returns bbox as [[x1,y1],[x2,y1],[x2,y2],[x1,y2]]
                    # Convert to x, y, width, height format
                    x1, y1 = bbox[0]
                    x2, y2 = bbox[2]
                    
                    bounding_boxes.append({
                        'text': text,
                        'confidence': int(confidence * 100),  # Convert to percentage
                        'x': int(x1),
                        'y': int(y1),
                        'width': int(x2 - x1),
                        'height': int(y2 - y1),
                        'engine': 'easyocr'
                    })
            
            result['ocr_text'] = ' '.join(text_parts)
            result['raw_text'] = result['ocr_text']
            result['bounding_boxes'] = bounding_boxes
            result['ocr_engine'] = 'easyocr'
            
            return result
            
        except Exception as e:
            logger.error(f"EasyOCR failed: {e}, falling back to Tesseract")
            return self._perform_tesseract_ocr(image, result)
    
    def _combine_ocr_results(self, tesseract_result: Dict, easyocr_result: Dict) -> Dict:
        """Combine results from both OCR engines."""
        combined_result = tesseract_result.copy()
        
        # Combine text results
        tesseract_text = tesseract_result.get('ocr_text', '')
        easyocr_text = easyocr_result.get('ocr_text', '')
        
        # Use the longer text result (usually more accurate)
        if len(easyocr_text) > len(tesseract_text):
            combined_result['ocr_text'] = easyocr_text
            combined_result['raw_text'] = easyocr_text
            combined_result['primary_engine'] = 'easyocr'
        else:
            combined_result['primary_engine'] = 'tesseract'
        
        # Combine bounding boxes from both engines
        combined_bboxes = tesseract_result.get('bounding_boxes', []) + easyocr_result.get('bounding_boxes', [])
        combined_result['bounding_boxes'] = combined_bboxes
        combined_result['ocr_engine'] = 'both'
        
        # Store alternative text for comparison
        combined_result['alternative_text'] = {
            'tesseract': tesseract_text,
            'easyocr': easyocr_text
        }
        
        return combined_result
    
    def _perform_llm_ocr(self, image: np.ndarray, result: Dict) -> Dict:
        """Perform OCR using LLM models."""
        try:
            # Determine the best task type based on image characteristics
            task_type = self._determine_ocr_task_type(image)
            
            # Process with LLM
            llm_result = llm_ocr_processor.process_with_fallback(
                image=image,
                task_type=task_type
            )
            
            # Update result dictionary
            result['ocr_text'] = llm_result['text']
            result['raw_text'] = llm_result['text']
            result['confidence'] = llm_result['confidence']
            result['bounding_boxes'] = llm_result.get('bounding_boxes', [])
            result['structured_data'] = llm_result.get('structured_data', {})
            result['ocr_engine'] = 'llm'
            result['llm_metadata'] = llm_result.get('metadata', {})
            
            return result
            
        except Exception as e:
            logger.error(f"LLM OCR processing failed: {e}")
            raise
    
    def _perform_traditional_ocr(self, image: np.ndarray, result: Dict, engine: str) -> Dict:
        """Perform traditional OCR using specified engine."""
        if engine == 'easyocr' and self.easyocr_reader:
            return self._perform_easyocr(image, result)
        elif engine == 'both' and self.easyocr_reader:
            # Use both engines and combine results
            tesseract_result = self._perform_tesseract_ocr(image, result.copy())
            easyocr_result = self._perform_easyocr(image, result.copy())
            return self._combine_ocr_results(tesseract_result, easyocr_result)
        else:
            # Default to Tesseract
            return self._perform_tesseract_ocr(image, result)
    
    def _combine_llm_traditional_results(self, llm_result: Dict, traditional_result: Dict) -> Dict:
        """Combine LLM and traditional OCR results."""
        combined_result = llm_result.copy()
        
        # Get text from both methods
        llm_text = llm_result.get('ocr_text', '')
        traditional_text = traditional_result.get('ocr_text', '')
        
        # Determine which text to use as primary based on length and confidence
        llm_confidence = llm_result.get('confidence', 0.5)
        
        if llm_confidence >= 0.7 and len(llm_text) > len(traditional_text) * 0.8:
            # Use LLM result as primary
            combined_result['primary_engine'] = 'llm'
            combined_result['alternative_text'] = {
                'traditional': traditional_text,
                'llm': llm_text
            }
        else:
            # Use traditional result as primary
            combined_result['ocr_text'] = traditional_text
            combined_result['raw_text'] = traditional_text
            combined_result['primary_engine'] = 'traditional'
            combined_result['alternative_text'] = {
                'traditional': traditional_text,
                'llm': llm_text
            }
        
        # Combine bounding boxes
        traditional_bboxes = traditional_result.get('bounding_boxes', [])
        llm_bboxes = llm_result.get('bounding_boxes', [])
        combined_result['bounding_boxes'] = traditional_bboxes + llm_bboxes
        
        # Keep LLM metadata for cost tracking
        combined_result['llm_metadata'] = llm_result.get('llm_metadata', {})
        combined_result['ocr_engine'] = 'llm+traditional'
        
        return combined_result
    
    def _determine_ocr_task_type(self, image: np.ndarray) -> OCRTaskType:
        """Determine the appropriate OCR task type based on image characteristics."""
        height, width = image.shape[:2]
        
        # Simple heuristics to determine task type
        # In a real implementation, this could be more sophisticated
        
        # If image is very wide, might be a table
        aspect_ratio = width / height
        if aspect_ratio > 2.0:
            return OCRTaskType.TABLE_EXTRACTION
        
        # If image is large and complex, use document analysis
        if width * height > 1000000:  # > 1 megapixel
            return OCRTaskType.DOCUMENT_ANALYSIS
        
        # Default to basic text extraction
        return OCRTaskType.BASIC_TEXT_EXTRACTION
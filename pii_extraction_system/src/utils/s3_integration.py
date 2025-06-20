"""
S3 Integration for Automated Phase 0 Dataset Creation

This module provides S3 connectivity for batch document processing and automated
labeled dataset creation without manual uploads.
"""

import boto3
import asyncio
import logging
import tempfile
import os
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import json
import base64
from io import BytesIO
from PIL import Image

# For PDF to image conversion
try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

# For password-protected documents
try:
    import msoffcrypto
    MSOFFCRYPTO_AVAILABLE = True
except ImportError:
    MSOFFCRYPTO_AVAILABLE = False

from core.config import settings
from llm.multimodal_llm_service import llm_service
from core.logging_config import get_logger

logger = get_logger(__name__)


def convert_document_to_images(file_content: bytes, file_extension: str, password: Optional[str] = None) -> List[str]:
    """Convert document to base64-encoded images for vision processing"""
    images = []
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
        temp_file.write(file_content)
        temp_file_path = temp_file.name
    
    try:
        # Handle password-protected files
        processed_file_path = temp_file_path
        
        if password and file_extension.lower() in ['.pdf', '.docx', '.xlsx']:
            # For password-protected files, try to decrypt first
            if MSOFFCRYPTO_AVAILABLE and file_extension.lower() in ['.docx', '.xlsx']:
                try:
                    with open(temp_file_path, 'rb') as f:
                        office_file = msoffcrypto.OfficeFile(f)
                        office_file.load_key(password=password)
                        
                        # Create decrypted file
                        decrypted_path = temp_file_path + '_decrypted' + file_extension
                        with open(decrypted_path, 'wb') as decrypted_f:
                            office_file.decrypt(decrypted_f)
                        
                        processed_file_path = decrypted_path
                except Exception as e:
                    logger.warning(f"Failed to decrypt Office document: {e}")
        
        # Convert based on file type
        if file_extension.lower() == '.pdf':
            if PDF2IMAGE_AVAILABLE:
                try:
                    # For password-protected PDFs, pdf2image can handle them directly
                    kwargs = {'dpi': 200, 'fmt': 'PNG'}
                    if password:
                        kwargs['userpw'] = password
                    
                    # Convert PDF pages to images
                    pdf_images = convert_from_path(
                        processed_file_path,
                        **kwargs
                    )
                    
                    for img in pdf_images:
                        # Convert PIL image to base64
                        buffer = BytesIO()
                        img.save(buffer, format='PNG')
                        img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                        images.append(img_base64)
                        
                except Exception as e:
                    logger.error(f"PDF to image conversion failed: {e}")
                    # Try without password if it failed
                    if password:
                        try:
                            logger.info("Retrying PDF conversion without password...")
                            pdf_images = convert_from_path(
                                processed_file_path,
                                dpi=200,
                                fmt='PNG'
                            )
                            for img in pdf_images:
                                buffer = BytesIO()
                                img.save(buffer, format='PNG')
                                img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                                images.append(img_base64)
                        except Exception as e2:
                            logger.error(f"PDF conversion retry failed: {e2}")
        
        elif file_extension.lower() in ['.docx', '.doc']:
            # For Word documents
            if file_extension.lower() == '.doc':
                # Old .doc format not directly supported
                logger.warning(f"Old .doc format detected. Consider converting to .docx for better support.")
                # Create a placeholder image with instructions
                text_content = (
                    f"Document: {os.path.basename(processed_file_path)}\n\n"
                    "⚠️ Old .doc format is not directly supported\n\n"
                    "This is a legacy Microsoft Word format that cannot be\n"
                    "processed directly. Please:\n\n"
                    "1. Open the document in Microsoft Word\n"
                    "2. Save it as .docx format (File → Save As → .docx)\n"
                    "3. Re-upload the .docx version\n\n"
                    "Alternatively, you can:\n"
                    "- Export the document as PDF\n"
                    "- Take screenshots of the document pages"
                )
                img = create_text_image(text_content)
                buffer = BytesIO()
                img.save(buffer, format='PNG')
                img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                images.append(img_base64)
            else:
                try:
                    # Try to import docx here to avoid import issues
                    from docx import Document
                    doc = Document(processed_file_path)
                    text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                    
                    # Create a simple text image
                    img = create_text_image(text_content)
                    buffer = BytesIO()
                    img.save(buffer, format='PNG')
                    img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                    images.append(img_base64)
                except Exception as e:
                    logger.error(f"Word document processing failed: {e}")
        
        elif file_extension.lower() in ['.xlsx', '.xls']:
            # For Excel files, create a simple representation
            try:
                import openpyxl
                workbook = openpyxl.load_workbook(processed_file_path)
                text_content = ''
                
                for sheet in workbook.worksheets:
                    text_content += f"Sheet: {sheet.title}\n"
                    for row in sheet.iter_rows(values_only=True):
                        row_text = '\t'.join([str(cell) if cell is not None else '' for cell in row])
                        text_content += row_text + '\n'
                    text_content += '\n'
                
                # Create a text image
                img = create_text_image(text_content)
                buffer = BytesIO()
                img.save(buffer, format='PNG')
                img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                images.append(img_base64)
            except Exception as e:
                logger.error(f"Excel document processing failed: {e}")
        
        elif file_extension.lower() == '.txt':
            # For text files, create a text image
            try:
                with open(processed_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text_content = f.read()
                
                img = create_text_image(text_content)
                buffer = BytesIO()
                img.save(buffer, format='PNG')
                img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                images.append(img_base64)
            except Exception as e:
                logger.error(f"Text file processing failed: {e}")
        
        elif file_extension.lower() in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            # Direct image processing
            try:
                img = Image.open(processed_file_path)
                if img.mode == 'RGBA':
                    # Convert RGBA to RGB
                    background = Image.new('RGB', img.size, (255, 255, 255))
                    background.paste(img, mask=img.split()[-1])
                    img = background
                
                buffer = BytesIO()
                img.save(buffer, format='PNG')
                img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                images.append(img_base64)
                
            except Exception as e:
                logger.error(f"Image processing failed: {e}")
        
        else:
            logger.warning(f"Unsupported file type for vision processing: {file_extension}")
    
    finally:
        # Clean up temporary files
        try:
            os.unlink(temp_file_path)
            if processed_file_path != temp_file_path:
                os.unlink(processed_file_path)
        except Exception as e:
            logger.warning(f"Failed to clean up temporary files: {e}")
    
    return images


def create_text_image(text: str, width: int = 800, font_size: int = 12) -> Image:
    """Create a PIL image from text content"""
    from PIL import ImageDraw, ImageFont
    
    # Split text into lines and estimate image height
    lines = text.split('\n')
    max_chars_per_line = width // (font_size // 2)
    
    # Wrap long lines
    wrapped_lines = []
    for line in lines:
        if len(line) <= max_chars_per_line:
            wrapped_lines.append(line)
        else:
            # Simple word wrapping
            words = line.split(' ')
            current_line = ''
            for word in words:
                if len(current_line + ' ' + word) <= max_chars_per_line:
                    current_line += (' ' + word) if current_line else word
                else:
                    if current_line:
                        wrapped_lines.append(current_line)
                    current_line = word
            if current_line:
                wrapped_lines.append(current_line)
    
    # Calculate image height
    line_height = font_size + 4
    height = max(len(wrapped_lines) * line_height + 40, 400)
    
    # Create image
    img = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(img)
    
    # Try to use a default font
    try:
        font = ImageFont.truetype('/System/Library/Fonts/Arial.ttf', font_size)
    except:
        font = ImageFont.load_default()
    
    # Draw text
    y_position = 20
    for line in wrapped_lines[:100]:  # Limit to first 100 lines
        draw.text((20, y_position), line, fill='black', font=font)
        y_position += line_height
    
    return img


@dataclass
class S3Document:
    """Represents a document in S3"""
    bucket: str
    key: str
    size: int
    last_modified: datetime
    
    @property
    def filename(self) -> str:
        return Path(self.key).name
    
    @property
    def extension(self) -> str:
        return Path(self.key).suffix.lower()


@dataclass
class DocumentProcessingResult:
    """Result of processing a single document"""
    s3_key: str
    filename: str
    success: bool
    entities: List[Dict] = None
    classification: Dict = None
    cost: float = 0.0
    processing_time: float = 0.0
    error: Optional[str] = None
    
    def to_phase0_format(self) -> Dict[str, Any]:
        """Convert to Phase 0 dataset format"""
        return {
            'id': f"s3_{hash(self.s3_key)}",
            'name': self.filename,
            'type': f"s3_document{Path(self.filename).suffix}",
            'size': 0,  # Not storing content locally
            'source': 's3',
            's3_key': self.s3_key,
            'uploaded_at': datetime.now().isoformat(),
            'labeled': self.success,
            'metadata': {
                'document_type': self.classification.get('domain_detail', 'Unknown') if self.classification else 'Unknown',
                'difficulty_level': self.classification.get('difficulty_level', 'Unknown') if self.classification else 'Unknown',
                'domain': self.classification.get('domain', 'Unknown') if self.classification else 'Unknown',
                'source': 's3',
                's3_bucket': self.s3_key.split('/')[0] if '/' in self.s3_key else ''
            },
            'gpt4o_labels': {
                'method': 'gpt4o_vision_s3' if self.success else 'failed',
                'entities': self.entities or [],
                'cost': self.cost,
                'processing_time': self.processing_time,
                'error': self.error
            } if self.entities or self.error else None,
            'validation_status': 'Auto-labeled' if self.success else 'Failed'
        }


@dataclass
class BatchProcessingResult:
    """Result of batch processing operation"""
    total_documents: int
    successful_documents: int
    failed_documents: int
    total_cost: float
    total_processing_time: float
    dataset_s3_key: Optional[str] = None
    error_summary: List[str] = None
    
    @property
    def success_rate(self) -> float:
        return self.successful_documents / max(self.total_documents, 1)


class S3DocumentProcessor:
    """Processes documents from S3 for Phase 0 dataset creation"""
    
    def __init__(self, 
                 bucket_name: str,
                 aws_region: str = "us-west-2",
                 aws_access_key: Optional[str] = None,
                 aws_secret_key: Optional[str] = None):
        """
        Initialize S3 processor
        
        Args:
            bucket_name: S3 bucket name
            aws_region: AWS region
            aws_access_key: AWS access key (optional, uses IAM if not provided)
            aws_secret_key: AWS secret key (optional, uses IAM if not provided)
        """
        self.bucket_name = bucket_name
        self.aws_region = aws_region
        
        # Initialize S3 client
        s3_kwargs = {'region_name': aws_region}
        if aws_access_key and aws_secret_key:
            s3_kwargs.update({
                'aws_access_key_id': aws_access_key,
                'aws_secret_access_key': aws_secret_key
            })
        
        self.s3_client = boto3.client('s3', **s3_kwargs)
        
        # Supported file extensions
        self.supported_extensions = {'.pdf', '.docx', '.xlsx', '.png', '.jpg', '.jpeg', '.tiff'}
        
        logger.info(f"S3 processor initialized for bucket: {bucket_name}")
    
    def discover_documents(self, 
                          prefix: str = "",
                          max_documents: Optional[int] = None,
                          skip_processed: bool = True) -> List[S3Document]:
        """
        Discover documents in S3 bucket
        
        Args:
            prefix: S3 prefix to search (e.g., "documents/")
            max_documents: Maximum number of documents to return
            skip_processed: Skip documents that have been processed
            
        Returns:
            List of S3Document objects
        """
        documents = []
        
        try:
            paginator = self.s3_client.get_paginator('list_objects_v2')
            page_iterator = paginator.paginate(
                Bucket=self.bucket_name,
                Prefix=prefix,
                MaxKeys=1000
            )
            
            for page in page_iterator:
                for obj in page.get('Contents', []):
                    key = obj['Key']
                    
                    # Check if it's a supported file type
                    if Path(key).suffix.lower() in self.supported_extensions:
                        
                        # Skip if already processed (check for .processed marker)
                        if skip_processed and self._is_already_processed(key):
                            continue
                        
                        # Check file size (skip if too large)
                        if obj['Size'] > settings.processing.max_file_size_mb * 1024 * 1024:
                            logger.warning(f"Skipping large file: {key} ({obj['Size']} bytes)")
                            continue
                        
                        documents.append(S3Document(
                            bucket=self.bucket_name,
                            key=key,
                            size=obj['Size'],
                            last_modified=obj['LastModified']
                        ))
                        
                        # Check document limit
                        if max_documents and len(documents) >= max_documents:
                            break
                
                if max_documents and len(documents) >= max_documents:
                    break
                    
        except Exception as e:
            logger.error(f"Error discovering S3 documents: {e}")
            raise
        
        logger.info(f"Discovered {len(documents)} documents in S3")
        return documents
    
    def _is_already_processed(self, document_key: str) -> bool:
        """Check if document has already been processed"""
        processed_marker_key = f"{document_key}.processed"
        
        try:
            self.s3_client.head_object(Bucket=self.bucket_name, Key=processed_marker_key)
            return True
        except self.s3_client.exceptions.NoSuchKey:
            return False
        except Exception:
            return False
    
    def _mark_as_processed(self, document_key: str, result: DocumentProcessingResult):
        """Mark document as processed with metadata"""
        processed_marker_key = f"{document_key}.processed"
        
        processing_metadata = {
            'processed_at': datetime.now().isoformat(),
            'success': result.success,
            'entities_found': len(result.entities) if result.entities else 0,
            'cost': result.cost,
            'processing_time': result.processing_time,
            'error': result.error
        }
        
        try:
            self.s3_client.put_object(
                Bucket=self.bucket_name,
                Key=processed_marker_key,
                Body=json.dumps(processing_metadata, indent=2),
                ContentType='application/json'
            )
        except Exception as e:
            logger.warning(f"Failed to mark {document_key} as processed: {e}")
    
    async def process_document(self, 
                             document: S3Document, 
                             password: str = "Hubert",
                             model_key: str = "openai/gpt-4o-mini") -> DocumentProcessingResult:
        """
        Process a single S3 document
        
        Args:
            document: S3Document to process
            password: Password for encrypted documents
            model_key: LLM model to use for processing
            
        Returns:
            DocumentProcessingResult
        """
        start_time = datetime.now()
        
        try:
            logger.debug(f"Processing S3 document: {document.key}")
            
            # Download document content
            response = self.s3_client.get_object(Bucket=document.bucket, Key=document.key)
            file_content = response['Body'].read()
            
            # Convert document to images for vision processing
            images = convert_document_to_images(file_content, document.extension, password)
            
            if not images:
                return DocumentProcessingResult(
                    s3_key=document.key,
                    filename=document.filename,
                    success=False,
                    error="Could not convert document to images"
                )
            
            # Process with LLM vision
            all_entities = []
            total_cost = 0.0
            classification = None
            
            for i, image_data in enumerate(images):
                result = llm_service.extract_pii_from_image(
                    image_data,
                    model_key,
                    document_type="business_document"
                )
                
                if result.get('success'):
                    # Add page information to entities
                    page_entities = result.get('pii_entities', [])
                    for entity in page_entities:
                        entity['page'] = i + 1
                        entity['source'] = f's3_vision_{model_key}'
                    
                    all_entities.extend(page_entities)
                    total_cost += result.get('usage', {}).get('estimated_cost', 0)
                    
                    # Get classification from first page
                    if i == 0 and result.get('structured_data'):
                        classification = result['structured_data'].get('document_classification')
                else:
                    logger.warning(f"Failed to process page {i+1} of {document.key}: {result.get('error')}")
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            # Create result
            result = DocumentProcessingResult(
                s3_key=document.key,
                filename=document.filename,
                success=True,
                entities=all_entities,
                classification=classification,
                cost=total_cost,
                processing_time=processing_time
            )
            
            # Mark as processed
            self._mark_as_processed(document.key, result)
            
            logger.info(f"Successfully processed {document.key}: {len(all_entities)} entities, ${total_cost:.4f}")
            return result
            
        except Exception as e:
            processing_time = (datetime.now() - start_time).total_seconds()
            error_msg = str(e)
            
            logger.error(f"Failed to process {document.key}: {error_msg}")
            
            result = DocumentProcessingResult(
                s3_key=document.key,
                filename=document.filename,
                success=False,
                error=error_msg,
                processing_time=processing_time
            )
            
            # Still mark as processed (with error) to avoid reprocessing
            self._mark_as_processed(document.key, result)
            
            return result
    
    async def process_batch(self, 
                          documents: List[S3Document],
                          password: str = "Hubert",
                          model_key: str = "openai/gpt-4o-mini",
                          max_concurrent: int = 5,
                          max_budget: float = 100.0) -> BatchProcessingResult:
        """
        Process a batch of S3 documents
        
        Args:
            documents: List of S3Documents to process
            password: Password for encrypted documents  
            model_key: LLM model to use
            max_concurrent: Maximum concurrent processing
            max_budget: Maximum budget for batch processing
            
        Returns:
            BatchProcessingResult
        """
        logger.info(f"Starting batch processing of {len(documents)} documents")
        
        # Semaphore to limit concurrent processing
        semaphore = asyncio.Semaphore(max_concurrent)
        
        async def process_with_semaphore(doc):
            async with semaphore:
                return await self.process_document(doc, password, model_key)
        
        # Process documents
        start_time = datetime.now()
        results = []
        total_cost = 0.0
        
        # Process in chunks to monitor budget
        chunk_size = 10
        for i in range(0, len(documents), chunk_size):
            chunk = documents[i:i + chunk_size]
            
            # Process chunk
            chunk_tasks = [process_with_semaphore(doc) for doc in chunk]
            chunk_results = await asyncio.gather(*chunk_tasks, return_exceptions=True)
            
            # Handle results and budget checking
            for result in chunk_results:
                if isinstance(result, Exception):
                    logger.error(f"Processing exception: {result}")
                    continue
                    
                results.append(result)
                total_cost += result.cost
                
                # Check budget limit
                if total_cost > max_budget:
                    logger.warning(f"Budget limit reached: ${total_cost:.2f} > ${max_budget}")
                    break
            
            if total_cost > max_budget:
                break
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        # Calculate summary statistics
        successful_results = [r for r in results if r.success]
        failed_results = [r for r in results if not r.success]
        
        batch_result = BatchProcessingResult(
            total_documents=len(results),
            successful_documents=len(successful_results),
            failed_documents=len(failed_results),
            total_cost=total_cost,
            total_processing_time=processing_time,
            error_summary=[r.error for r in failed_results if r.error]
        )
        
        logger.info(f"Batch processing complete: {batch_result.successful_documents}/{batch_result.total_documents} successful, ${total_cost:.2f} cost")
        
        return batch_result, results
    
    def export_dataset_to_s3(self, 
                           processing_results: List[DocumentProcessingResult],
                           output_prefix: str = "labeled-datasets/") -> str:
        """
        Export processed results as Phase 0 dataset to S3
        
        Args:
            processing_results: List of processing results
            output_prefix: S3 prefix for output dataset
            
        Returns:
            S3 key of exported dataset
        """
        # Convert results to Phase 0 format
        phase0_dataset = []
        for result in processing_results:
            if result.success:
                phase0_dataset.append(result.to_phase0_format())
        
        # Create dataset metadata
        dataset_metadata = {
            'created_at': datetime.now().isoformat(),
            'total_documents': len(processing_results),
            'successful_documents': len(phase0_dataset),
            'total_entities': sum(len(doc['gpt4o_labels']['entities']) for doc in phase0_dataset),
            'total_cost': sum(result.cost for result in processing_results),
            'source': 's3_batch_processing',
            'bucket': self.bucket_name
        }
        
        # Create export data
        export_data = {
            'metadata': dataset_metadata,
            'documents': phase0_dataset
        }
        
        # Export to S3
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_key = f"{output_prefix}phase0_dataset_{timestamp}.json"
        
        try:
            self.s3_client.put_object(
                Bucket=self.bucket_name,
                Key=output_key,
                Body=json.dumps(export_data, indent=2, default=str),
                ContentType='application/json'
            )
            
            logger.info(f"Dataset exported to S3: s3://{self.bucket_name}/{output_key}")
            return output_key
            
        except Exception as e:
            logger.error(f"Failed to export dataset to S3: {e}")
            raise


# Convenience functions for integration

def create_s3_processor_from_config() -> S3DocumentProcessor:
    """Create S3 processor from configuration"""
    if settings.data_source.source_type != 's3':
        raise ValueError("Data source not configured for S3")
    
    if not settings.data_source.s3_bucket:
        raise ValueError("S3 bucket not configured")
    
    return S3DocumentProcessor(
        bucket_name=settings.data_source.s3_bucket,
        aws_region=settings.data_source.s3_region,
        aws_access_key=settings.aws_access_key_id,
        aws_secret_key=settings.aws_secret_access_key
    )


async def process_s3_bucket_to_phase0(bucket_name: str,
                                    prefix: str = "",
                                    password: str = "Hubert",
                                    max_documents: int = 100,
                                    max_budget: float = 50.0) -> Dict[str, Any]:
    """
    High-level function to process S3 bucket and create Phase 0 dataset
    
    Args:
        bucket_name: S3 bucket name
        prefix: S3 prefix to search
        password: Password for encrypted documents
        max_documents: Maximum documents to process
        max_budget: Maximum budget for processing
        
    Returns:
        Dictionary with processing results and dataset location
    """
    # Initialize processor
    processor = S3DocumentProcessor(bucket_name)
    
    # Discover documents
    documents = processor.discover_documents(prefix=prefix, max_documents=max_documents)
    
    if not documents:
        return {
            'success': False,
            'error': 'No documents found in S3 bucket',
            'documents_found': 0
        }
    
    # Process batch
    batch_result, processing_results = await processor.process_batch(
        documents, 
        password=password,
        max_budget=max_budget
    )
    
    # Export dataset
    dataset_key = None
    if processing_results:
        dataset_key = processor.export_dataset_to_s3(processing_results)
    
    return {
        'success': True,
        'batch_result': batch_result,
        'dataset_s3_key': dataset_key,
        'documents_found': len(documents),
        'documents_processed': batch_result.total_documents,
        'success_rate': batch_result.success_rate,
        'total_cost': batch_result.total_cost
    }